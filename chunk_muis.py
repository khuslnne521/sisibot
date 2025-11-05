# -*- coding: utf-8 -*-
"""
DOCX -> chunks.jsonl (RAG-д бэлдэх)
- Нэг фолдер (option: рекурсив) доторх бүх .docx файлыг уншина
- OCR artifacts цэвэрлэнэ
- Гарчгаар (heading) секцлээд өгүүлбэрийн цонхоор chunk хийнэ (overlap-той)
- Нэгтгэсэн болон файл тус бүрийн JSONL бичнэ

Шаардлага:
  pip install python-docx tiktoken

Ашиглах жишээ:
  python chunk_muis.py \
    --input-dir "/Users/macbookpro/iCloud Drive (Archive)/Documents/muisDuremJuram/muisDuremOCR" \
    --out "muis_chunks.jsonl" \
    --max-tokens 600 --overlap 80 --recursive
"""

import os, re, json, uuid, argparse, sys
from datetime import datetime
from typing import List

# ========== Токен тоолох ==========
def get_token_counter():
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        def count_tokens(text: str) -> int:
            return len(enc.encode(text))
        return count_tokens, "token"
    except Exception:
        def count_tokens(text: str) -> int:
            return max(1, len(text) // 4)  # ~4 тэмдэгт ≈ 1 токен
        return count_tokens, "approx_char/4"

count_tokens, token_mode = get_token_counter()

# ========== OCR цэвэрлэгээ ==========
CLEAN_PATTERNS = [
    (r"[ \t]+", " "),         # олон зайг нэг болгох
    (r"\s+\n", "\n"),         # мөрийн төгсгөлийн хоосон зай
    (r"\n{3,}", "\n\n"),      # 3+ хоосон мөрийг 2 болгох
    (r"[“”]", '"'), (r"[‘’]", "'"),
    (r"\u200b", ""),           # zero-width
    (r"[|]{2,}", "|"),
    (r"[^\S\r\n]{2,}", " "),
]

def clean_text(s: str) -> str:
    s = s.strip()
    for pat, rep in CLEAN_PATTERNS:
        s = re.sub(pat, rep, s)
    s = re.sub(r"(\.){3,}", "…", s)  # ...
    return s

# ========== Өгүүлбэрээр хуваах ==========
# Төгсгөл тэмдэг (., ?, !, …) + дараагийн өгүүлбэрийн эхлэл (Монгол/Англи том үсэг, тоо, ишлэл)
SENT_SPLIT = re.compile(r"(?<=[\.\?\!…])\s+(?=[А-ЯA-Z0-9“\"'])")

def split_sentences(text: str):
    """Текстийг өгүүлбэрүүдэд хуваана (богино бол бүхэлд нь буцаана)."""
    text = clean_text(text)
    if len(text) < 50:
        return [text]
    parts = SENT_SPLIT.split(text)
    return [p.strip() for p in parts if p and p.strip()]

# ========== .docx унших (гарчиг барих гэж оролдох) ==========
def read_docx_with_headings(path: str):
    """
    python-docx ашиглан DOCX уншаад paragraphs-ийг гүйлгэнэ.
    Paragraph.style.name-аас Heading 1..n (эсвэл 'Гарчиг 1..n')-г илрүүлж
    шинэ блок эхлүүлнэ. Үгүй бол одоогийн блокт текст нэмж явна.
    """
    from docx import Document
    doc = Document(path)
    blocks = []
    current = {"heading_levels": [], "texts": []}

    def flush():
        nonlocal current, blocks
        if current["heading_levels"] or current["texts"]:
            blocks.append(current)
            current = {"heading_levels": [], "texts": []}

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = getattr(p.style, "name", "") or ""
        style_lc = style.lower()

        # Heading 1..n, эсвэл "Гарчиг 1..n" маягийн стил
        m = re.match(r"heading\s*(\d+)", style_lc) or re.match(r"гарчиг\s*(\d+)", style_lc)
        if m:
            flush()
            level = int(m.group(1))
            current = {"heading_levels": [(level, text)], "texts": []}
        else:
            current["texts"].append(text)

    flush()
    if not blocks:
        # heading олдоогүй бол бүх текстийг нэг блокт
        full = "\n".join([p.text for p in doc.paragraphs if (p.text or "").strip()])
        blocks = [{"heading_levels": [], "texts": [full]}]
    return blocks

# ========== Секц бүрийг токен цонхоор chunk хийх ==========
def make_chunks_from_section(section_text: str, max_tokens=600, overlap_tokens=80) -> List[str]:
    """
    Өгүүлбэрүүдийг цуглуулж яваад max_tokens давахад chunk тасална.
    Дараагийн chunk эхлэхдээ өмнөхөөс overlap_tokens орчмыг давхардуулж авна.
    """
    sents = split_sentences(section_text)
    chunks, buff, buff_tok = [], [], 0
    for sent in sents:
        tks = count_tokens(sent)
        if buff_tok + tks > max_tokens and buff:
            chunks.append("\n".join(buff).strip())
            # overlap
            overlap, tok_sum = [], 0
            for s in reversed(buff):
                tok_sum += count_tokens(s)
                overlap.append(s)
                if tok_sum >= overlap_tokens:
                    break
            buff = list(reversed(overlap))
            buff_tok = sum(count_tokens(x) for x in buff)
        buff.append(sent)
        buff_tok += tks
    if buff:
        chunks.append("\n".join(buff).strip())
    return chunks

def chunk_docx_file(
    path: str,
    source_id: str = None,
    max_tokens: int = 600,
    overlap_tokens: int = 80,
    h_prefix_levels: int = 3
):
    
    """
    Нэг файл дээрх бүх блокуудыг chunk хийгээд JSONL мөрүүдийн жагсаалт буцаана.
    heading_path нь эхний h_prefix_levels хүртэлх гарчгуудыг “A > B > …” байдлаар.
    """
    source_id = source_id or os.path.basename(path)
    blocks = read_docx_with_headings(path)
    out = []
    section_idx = 0
    for blk in blocks:
        section_idx += 1
        heading_path = []
        for lvl, txt in blk.get("heading_levels", []):
            if lvl <= h_prefix_levels:
                heading_path.append(clean_text(txt))
        body = clean_text("\n".join(blk.get("texts", [])))
        if not body:
            continue
        chunks = make_chunks_from_section(body, max_tokens=max_tokens, overlap_tokens=overlap_tokens)
        for i, ch in enumerate(chunks, start=1):
            preview = re.sub(r"\s+", " ", ch)[:160]
            out.append({
                "id": str(uuid.uuid4()),
                "source_id": source_id,
                "section_index": section_idx,
                "chunk_index": i,
                "heading_path": heading_path,
                "token_estimate": count_tokens(ch),
                "token_mode": token_mode,
                "text": ch,
                "preview": preview
            })
    return out

# ========== Фолдер скан хийх ==========
def find_docx_files(root_dir: str, recursive: bool = True) -> List[str]:

    """root_dir дахь .docx файлуудыг (рекурсив бол бүх дэд хавтас) олж жагсаана."""
    docx_files = []
    if recursive:
        for base, _, files in os.walk(root_dir):
            for fn in files:
                if fn.lower().endswith(".docx"):
                    docx_files.append(os.path.join(base, fn))
    else:
        for fn in os.listdir(root_dir):
            p = os.path.join(root_dir, fn)
            if os.path.isfile(p) and fn.lower().endswith(".docx"):
                docx_files.append(p)
    docx_files.sort()
    return docx_files

def write_jsonl(path: str, rows: List[dict]):
    """Мөр бүрийг JSONL хэлбэрээр бичих энгийн writer."""
    with open(path, "w", encoding="utf-8") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")

def main():
    parser = argparse.ArgumentParser(description="Chunk DOCX directory for RAG")
    parser.add_argument("--input-dir", required=True, help="DOCX файлууд байрлах фолдер")
    parser.add_argument("--out", default=None, help="Нэгтгэсэн JSONL гаралт (default: chunks_<ts>.jsonl)")
    parser.add_argument("--max-tokens", type=int, default=600)
    parser.add_argument("--overlap", type=int, default=80)
    parser.add_argument("--recursive", action="store_true", help="Дотогш рекурсив хайх")
    parser.add_argument("--per-file-jsonl", action="store_true", help="Файл тус бүрийн chunks_<filename>.jsonl бичих")
    args = parser.parse_args()

    indir = args.input_dir
    if not os.path.isdir(indir):
        print(f"[ERROR] Фолдер олдсонгүй: {indir}", file=sys.stderr)
        sys.exit(1)

    files = find_docx_files(indir, recursive=args.recursive)
    if not files:
        print(f"[WARN] .docx файл олдсонгүй: {indir}")
        sys.exit(0)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_path = args.out or f"chunks_{ts}.jsonl"
    all_chunks = []
    print(f"[INFO] Нийт {len(files)} .docx файл олдлоо.")
    for idx, fp in enumerate(files, start=1):
        try:
            chunks = chunk_docx_file(
                fp,
                source_id=os.path.relpath(fp, indir),  # харьцангуй замыг source_id болгох
                max_tokens=args.max_tokens,
                overlap_tokens=args.overlap
            )
            all_chunks.extend(chunks)
            print(f"[{idx}/{len(files)}] OK - {os.path.basename(fp)} -> {len(chunks)} chunks")
            if args.per_file_jsonl:
                base = os.path.splitext(os.path.basename(fp))[0]
                per_path = f"chunks_{base}_{ts}.jsonl"
                write_jsonl(per_path, chunks)
        except Exception as e:
            print(f"[{idx}/{len(files)}] FAIL - {os.path.basename(fp)} :: {e}", file=sys.stderr)

    write_jsonl(out_path, all_chunks)
    print(f"==> Нийт {len(all_chunks)} chunk бичлээ: {out_path}")

if __name__ == "__main__":
    main()
